import logging
import json

from django.conf import settings
import os
import docx

logger = logging.getLogger(__name__)


def export(dir_name: str, file_path: str, target: str, output_format: str, merged_format: bool) -> str:
    """
    Export the user edited output file to the desired format.
    :param dir_name: transcription directory
    :param target: notes or edited output
    :param file_path: The path to the file that should be exported.
    :param output_format: The export format.
    :param merged_format: Turn on merge of segments with same speaker.
    :return: Path to exported file.
    """
    if target == 'notes':
        return _export_notes(dir_name, file_path, output_format)
    else:
        return _export_edited_output(dir_name, file_path, output_format, merged_format)


def _export_notes(dir_name: str, file_path: str, output_format: str):
    if output_format == "json":
        # no processing needed, just return file path
        return os.path.join(settings.MEDIA_ROOT, dir_name, 'data', 'notes.json')
    data = {}
    try:
        with open(file_path, 'r') as f:
            data = json.load(f)
    except Exception as e:
        logger.error(f"Error reading data file for export: '{file_path}': {e}")

    if output_format == "docx":
        return _export_notes_docx(dir_name, data)
    elif output_format == "txt":
        return _export_notes_txt(dir_name, data)
    else:
        logger.error(f"Unknown export format: {output_format}")
        return ""


def _export_edited_output(dir_name: str, file_path: str, output_format: str, merged_format: bool) -> str:
    if output_format == "json" and not merged_format:
        # no processing needed, just return file path
        return os.path.join(settings.MEDIA_ROOT, dir_name, 'data', 'edited_output.json')
    data = {}
    if merged_format:
        data = _merge_speakers(file_path)
    else:
        try:
            with open(file_path, 'r') as f:
                data = json.load(f)
        except Exception as e:
            logger.error(f"Error reading data file for export: '{file_path}': {e}")

    if output_format == "json":
        output_file_path = os.path.join(settings.MEDIA_ROOT, dir_name, 'data', 'edited_output_merged.json')
        try:
            with open(output_file_path, 'w') as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            logger.error(f"Error writing edited output JSON with merged speaker format: '{output_file_path}': {e}")
        return output_file_path
    elif output_format == "docx":
        return _export_edited_output_docx(dir_name, data)
    elif output_format == "txt":
        return _export_edited_output_txt(dir_name, data)
    else:
        logger.error(f"Unknown export format: {output_format}")
        return ""


def _export_edited_output_docx(dir_name: str, data: dict):
    output_file_path = os.path.join(settings.MEDIA_ROOT, dir_name, 'data', 'edited_output.docx')
    try:
        document = docx.Document()
        p = document.add_paragraph()
        for line in data["lines"]:
            time = p.add_run(f"{line["startTime"]} - {line["endTime"]}")
            time.italic = True
            p.add_run("\t")
            p.add_run(f'{extract_speaker(line).strip()}\n')
            p.add_run("\t")
            p.add_run(f'{line["text"].strip()}\n')
        document.save(output_file_path)
    except Exception as e:
        logger.error(f"Error exporting word document: {e}")
        return ""
    return output_file_path

def _export_edited_output_txt(dir_name: str, data: dict):
    output_file_path = os.path.join(settings.MEDIA_ROOT, dir_name, 'data', 'edited_output.txt')
    try:
        output_lines = []
        for line in data["lines"]:
            speaker = extract_speaker(line).strip()
            output_lines.append(f"{line["startTime"]} - {line["endTime"]} \t {speaker}\n")
            output_lines.append(f"\t {line["text"]}\n")
        with open(output_file_path, 'w') as f:
            f.writelines(output_lines)
    except Exception as e:
        logger.error(f"Error exporting txt document: {e}")
        return ""
    return output_file_path


def _export_notes_docx(dir_name: str, data: dict):
    output_file_path = os.path.join(settings.MEDIA_ROOT, dir_name, 'data', 'notes.docx')
    try:
        document = docx.Document()
        p = document.add_paragraph()
        for line in data["notes"]:
            p.add_run(f"{line["date"]}\n")
            p.add_run(f"{line["note"]}\n\n")
        document.save(output_file_path)
    except Exception as e:
        logger.error(f"Error exporting notes in word document: {e}")
        return ""
    return output_file_path


def _export_notes_txt(dir_name: str, data: dict):
    output_file_path = os.path.join(settings.MEDIA_ROOT, dir_name, 'data', 'notes.txt')
    try:
        output_lines = []
        for line in data["notes"]:
            output_lines.append(f"{line["date"]}\n")
            output_lines.append(f"{line["note"]}\n\n")
        with open(output_file_path, 'w') as f:
            f.writelines(output_lines)
    except Exception as e:
        logger.error(f"Error exporting notes in txt format: {e}")
        return ""
    return output_file_path


def _merge_speakers(edit_output_path: str) -> dict:
    merged_data = {"lines": []}
    try:
        with open(edit_output_path, 'r') as f:
            data = json.load(f)

        current_speaker: str = ""
        for i, line in enumerate(data.get('lines', [])):
            if i==0:
                # add the first line
                new_line = {
                    "id": 0,
                    "startTime": line["startTime"],
                    "endTime": line["endTime"],
                    "speakerDesignation": line["speakerDesignation"],
                    "text": line["text"]
                }
                merged_data["lines"].append(new_line)
                current_speaker = line["speakerDesignation"]
            else:
                speaker = extract_speaker(line)
                if speaker != current_speaker:
                    # add new line
                    new_line = {
                        "id": i,
                        "startTime": line["startTime"],
                        "endTime": line["endTime"],
                        "speakerDesignation": speaker,
                        "text": line["text"],
                    }
                    merged_data["lines"].append(new_line)
                    current_speaker = speaker
                else:
                    # add the text to the previous line
                    merged_data["lines"][-1]["text"] = (merged_data["lines"][-1]["text"].strip()
                                                                      + " "
                                                                      + line["text"].strip())
                    # push the end time forward
                    merged_data["lines"][-1]["endTime"] = line["endTime"]

    except Exception as e:
        logger.error(f"Error when generating merged speaker format for output file '{edit_output_path}': {e}")

    return merged_data


def extract_speaker(line) -> str:
    if "speakerDesignation" in line:
        return line["speakerDesignation"]
    else:
        return "Undetermined speaker"
